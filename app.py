import os
import re
import random
import smtplib
import json
from datetime import datetime
from flask import Flask, render_template, request, redirect, flash, send_from_directory
from email.message import EmailMessage
from werkzeug.utils import secure_filename
from google.oauth2.service_account import Credentials
import gspread
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
load_dotenv()


UPLOAD_FOLDER = 'static/uploads'
TEMPLATE_PATH = 'Template.docx'
CREDENTIALS_PATH = 'ecx-progression-app-460516-54ee54fb1563.json'
IT_CREDENTIALS_PATH = 'it-progression-form-83b2a28b0704.json'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}
SCOPE = ["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"]
#------------------------------------------------ LIVE -----------------------------------------------------------------
# Email
SENDER_EMAIL = os.getenv("ECX_SENDER_EMAIL")
EMAIL_PASSWORD = os.getenv("ECX_EMAIL_PASSWORD")

# Google Sheets setup
creds_dict = json.loads(os.getenv("GOOGLE_CREDS_JSON"))
creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPE)
client = gspread.authorize(creds)
sheet = client.open("ECX PROGRESSION").sheet1

# IT Form Google Sheets Setup
it_creds_dict = json.loads(os.getenv("GOOGLE_IT_CREDS_JSON"))
it_creds = Credentials.from_service_account_info(it_creds_dict, scopes=SCOPE)
it_client = gspread.authorize(it_creds)
it_sheet = it_client.open("IT REQUEST FORM REPOSITORY").sheet1

app = Flask(__name__)
app.secret_key = "ecx-secret"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

#------------------------------------------------ LOCAL TESTING --------------------------------------------------------
# # Email
# SENDER_EMAIL = "ecxoperationalcompliance@gmail.com"
# EMAIL_PASSWORD = "bverlbfblogutkkf"
#
# # Google Sheets setup
# creds = Credentials.from_service_account_file(CREDENTIALS_PATH, scopes=SCOPE)
# client = gspread.authorize(creds)
# sheet = client.open("ECX PROGRESSION").sheet1
#
# # IT Form Google Sheets Setup
# it_creds = Credentials.from_service_account_file(IT_CREDENTIALS_PATH, scopes=SCOPE)
# it_client = gspread.authorize(it_creds)
# it_sheet = it_client.open("IT REQUEST FORM REPOSITORY").sheet1
#
# app = Flask(__name__)
# app.secret_key = "ecx-secret"
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

#----------------------------------------------- START OF WEB-APP ------------------------------------------------------

FIELDS = [
    "Employee Email", "Immediate Supervisor Email", "Reported By", "Reported By Title/Role", "Date of Report",
    "Employee Name", "Employee Title/Role", "Date of Incident", "Time of Incident",
    "Immediate Supervisor", "Department Head", "Alleged Violation", "Location",
    "Specific Area of Location", "Additional Person(s) Involved", "Witnesses",
    "Incident Description", "Employee Explanation", "Action Taken by Immediate Supervisor",
    "Immediate Supervisor Recommendation"
]

RECOMMENDATION_OPTIONS = [
    "Select recommendation...", "NA", "Coaching", "Verbal Warning",
    "1st Written Warning", "2nd Written Warning", "3rd Written Warning",
    "Final Warning", "For HR Review", "For HR Intervention", "For Management Review"
]


MULTILINE_FIELDS = {
    "Additional Person(s) Involved", "Witnesses",
    "Incident Description", "Employee Explanation",
    "Action Taken by Immediate Supervisor"
}

DEPARTMENT_HEAD_OPTIONS = [
    "Select Department Head...","Amelyn Talastas", "Monique Curia", "Samuel Pastrana", "Jaycee Quijano"
]

DEPARTMENT_HEAD_EMAILS = {
    "Amelyn Talastas": "amelyntalastas.ecx@gmail.com",
    "Monique Curia": "moniquecuria.ecx@gmail.com",
    "Samuel Pastrana": "samuelpastrana.ecx@gmail.com",
    "Jaycee Quijano": "jaycee@ecxperience.com"
}

ROLE_OPTIONS = [
    "Select Title/Role...", "CSA Agent", "Vouchering Agent", "OTR Agent", "QA AGENT", "Accounting Agent", "Team Lead","Trainer", "Data Scientist", "Assistant Operations Manager",
    "Operations Manager","Operations Success Coordinator"
]

ALLEGED_OPTIONS = [
    "Select Allegations...", "N/A", "Wasting time or loitering on Company time. - Level 1", "Absence without permission or without reasonable cause. - Level 1",
    "Leaving work assignment during work hours without previous permission and without reasonable cause. - Level 2", "Leaving or abandoning work assignment during official working hours. - Level 2",
    "Failure to render overtime work without a valid reason after signifying willingness to perform authorized overtime work. - Level 2", "Engaging in horse-play, officiousness and noisy conduct disturbing the work of other employees. - Level 2",
    "Doing private work during working hours without permission; Selling any kind of articles or lottery tickets, within the premises, without authorization of management. - Level 2",
    "Sleeping while on duty. - Level 2", "Willful or negligent disregard for standard operating procedures or processes. - Level 2", "Malingering or feigning illness to avoid doing assigned work during work hours. - Level 2",
    "Use of internet during office hours not related to work functions in the production area. - Level 2", "Instigating or willful disruption/ sabotage or slow-down of work. - Level 2",
    "Abuse of personal privileges such as extended breaks. - Level 2", "Bringing food and eating in production area (only hard candies and drinks in spill-proof canisters). - Level 2",
    "Unexcused absences or tardiness without prior notification to supervisor and Human Resources. For emergency leaves, unexcused absences or tardiness without notification to supervisor and Human Resources within the period provided in this handbook. - Level 2",
    "Going on under time without prior approval from the Immediate Supervisor/Manager. - Level 3", "Altering/manipulating timekeeping records - for a co-employee or for one's self. - Level 4"
]

def is_valid_email(email):
    return re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email)

def generate_incident_no():
    existing_ids = sheet.col_values(2)[1:]
    while True:
        ir = f"IR-{random.randint(1000000000, 9999999999)}"
        if ir not in existing_ids:
            return ir

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/")
def home():
    return render_template("Index.html")

@app.route("/progression_form", methods=["GET", "POST"])
def progression_form():
    if request.method == "POST":
        values = []
        errors = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        incident_no = generate_incident_no()
        values.extend([timestamp, incident_no])

        previous_values = {}

        for field in FIELDS:
            val = request.form.get(field, "").strip()
            previous_values[field] = val  # Store all user input
            if not val:
                errors.append(f"Please complete the '{field}' field.")
            values.append(val)

        # If there are errors, flash all of them and re-render the form
        if errors:
            for error in errors:
                flash(error, "danger")
            return render_template(
                "progression_form.html",
                fields=FIELDS,
                multiline_fields=MULTILINE_FIELDS,
                today=datetime.now().strftime("%Y-%m-%d"),
                recommendation_options=RECOMMENDATION_OPTIONS,
                department_head_options=DEPARTMENT_HEAD_OPTIONS,
                role_options=ROLE_OPTIONS,
                alleged_options=ALLEGED_OPTIONS,
                previous_values=previous_values  # So fields retain user input
            )

        # Email validation
        if not is_valid_email(values[2]) or not is_valid_email(values[3]):
            flash("Invalid email address.", "danger")
            today_str = datetime.now().strftime("%Y-%m-%d")
            return render_template(
                "progression_form.html",
                fields=FIELDS,
                multiline_fields=MULTILINE_FIELDS,
                generated_file=None,
                today=today_str,
                recommendation_options=RECOMMENDATION_OPTIONS,
                department_head_options=DEPARTMENT_HEAD_OPTIONS,
                role_options=ROLE_OPTIONS,
                alleged_options=ALLEGED_OPTIONS,
                previous_values=request.form
            )

        # Save to Google Sheet
        try:
            sheet.append_row(values)
        except Exception as e:
            flash("An error occurred while submitting to the spreadsheet. Please try again later.", "danger")
            print(f"[GSheet Error] {e}")
            return redirect(request.url)

        # Process uploaded images
        images = request.files.getlist("images")
        image_paths = []
        for image in images:
            if image and allowed_file(image.filename):
                filename = secure_filename(image.filename)
                save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image.save(save_path)
                image_paths.append(save_path)

        # Generate docx
        docx_path = os.path.join("generated", f"{incident_no}.docx")
        os.makedirs("generated", exist_ok=True)
        fill_docx(values, image_paths, docx_path, incident_no)

        # Email
        # Determine recipients
        recipients = [values[2], values[3], "manzonmarkjerwin@gmail.com"]
        dep_head = values[12]

        if dep_head in DEPARTMENT_HEAD_EMAILS:
            recipients.append(DEPARTMENT_HEAD_EMAILS[dep_head])

        # Eliminate duplicates (like repeated HR)
        recipients = list(set(email.lower() for email in recipients))

        send_email_with_attachment(
            recipients,
            f"Automated Notification: Incident Report Logged – {incident_no}",
            f"""This is an automated notification to inform you that an incident report has been logged in to the system:

{incident_no}

The matter is currently under review. Updates will be provided as they become available.

For questions or follow-ups, please contact our HR team at hr@ecxperience.com.

This is a system-generated message. No action is required unless otherwise indicated.""",
            docx_path
        )

        flash("Incident report submitted successfully.", "success")
        return redirect(f"/progression_form?generated={incident_no}")

    # GET method (default view)
    today_str = datetime.now().strftime("%Y-%m-%d")
    generated_file = request.args.get("generated")
    return render_template(
        "progression_form.html",
        fields=FIELDS,
        multiline_fields=MULTILINE_FIELDS,
        generated_file=generated_file,
        today=today_str,
        recommendation_options=RECOMMENDATION_OPTIONS,
        department_head_options=DEPARTMENT_HEAD_OPTIONS,
        role_options=ROLE_OPTIONS,
        alleged_options=ALLEGED_OPTIONS
    )


def fill_docx(data, images, save_path, incident_no):
    from docx.oxml import OxmlElement
    from docx.shared import Inches
    from docx import Document
    from datetime import datetime

    def insert_paragraph_after(paragraph):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        return paragraph._parent.add_paragraph()

    doc = Document(TEMPLATE_PATH)

    # Format Time of Incident
    try:
        raw_time = data[10]
        time_12hr = datetime.strptime(raw_time, "%H:%M").strftime("%I:%M %p")
    except Exception:
        time_12hr = data[10]

    placeholders = {
        '[Incident No.]': incident_no,
        '[Reported By]': data[4],
        '[Title / Role]': data[5],
        '[Date of Report]': data[6],
        '[Employee Name]': data[7],
        '[Employee Title / Role]': data[8],
        '[Date of Incident]': data[9],
        '[Time of Incident]': time_12hr,
        '[Immediate Supervisor]': data[11],
        '[Department Head]': data[12],
        '[Alleged Violation]': data[13],
        '[Location]': data[14],
        '[Specific Area of Location]': data[15],
        '[Additional Person(s) Involved]': data[16],
        '[Witnesses]': data[17],
        '[Incident Description]': data[18],
        '[Employee Explanation]': data[19],
        '[Action Taken]': data[20],
        '[Recommendation]': data[21],
    }

    def replace_in_paragraph(paragraph):
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, val in placeholders.items():
            new_text = new_text.replace(key, val)
        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ''
            paragraph.runs[0].text = new_text

    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    # Find the paragraph with "INDEX:"
    index_paragraph = None
    for para in doc.paragraphs:
        if para.text.strip().upper() == "INDEX:":
            index_paragraph = para
            break

    if index_paragraph:
        for img_path in images:
            run = index_paragraph.add_run()
            run.add_break()  # spacing after "INDEX:"
            run.add_picture(img_path, width=Inches(4))
            run.add_break()  # spacing between pictures

    else:
        for img_path in images:
            doc.add_picture(img_path, width=Inches(4))
            doc.add_paragraph("")

    doc.save(save_path)

@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory("generated", filename, as_attachment=True)

def send_email_with_attachment(to_emails, subject, body, attachment_path):
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = SENDER_EMAIL
    msg['To'] = ', '.join(to_emails)
    msg.set_content(body)

    with open(attachment_path, 'rb') as f:
        msg.add_attachment(f.read(), maintype='application', subtype='octet-stream', filename=os.path.basename(attachment_path))

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
            smtp.starttls()
            smtp.login(SENDER_EMAIL, EMAIL_PASSWORD)
            smtp.send_message(msg)
    except Exception as e:
        print(f"Failed to send email: {e}")
        flash("Form submitted, but failed to send email notification.", "danger")

#-------------------------------------------------- IT FORM -------------------------------------------------------------------------

def generate_ticket_no():
    existing_ids = sheet.col_values(2)[1:]
    while True:
        ir = random.randint(1000000000, 9999999999)
        if ir not in existing_ids:
            return ir

def send_email(to_emails, subject, body):
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = SENDER_EMAIL
    msg['To'] = ', '.join(to_emails)
    msg.set_content(body)

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
            smtp.starttls()
            smtp.login(SENDER_EMAIL, EMAIL_PASSWORD)
            smtp.send_message(msg)
    except Exception as e:
        print(f"Failed to send email: {e}")
        flash("Form submitted, but failed to send email notification.", "danger")

#------------------------------------------IT FORM GENERATION -----------------------------------------------------------------------

IT_FIELDS =["Email Address", "Agent Name", "Team Lead", "LOB", "Location", "PC Number", "Request Type", "Specific Request", "Date Requested"]

LOCATION_OPTIONS = ["Select Location...", "EC Cafe", "Eskina", "Parking Area", "Resto Bar", "Snack Bar", "Training Room", "Conference Room", "2nd Floor (COO Office)", "2nd Floor (HR Office)",
            "2nd Floor (Finance Office)", "3rd Floor (Phase 1)", "3rd Floor (Phase 2)", "4th Floor (Phase 1)", "4th Floor (Phase 2)", "5th Floor (Phase 1)", "5th Floor (CEO Office)"
            ]
REQUEST_OPTIONS = ["Select a Request...", "Assistance", "Network Problem", "Software Installation","System Issue", "Replacement", "Equipment Request"]

IT_MULTILINE_FIELDS = ["Specific Request"]

LOB_OPTIONS = ["Select LOB...", "N/A", "Accounting", "Banking Team", "Best Buy", "Buffer", "Convenience", "Dollar General", "Hospitality", "Industrial Retail", "M&I", "OTR", "Properties","QA", "Restaurant", "Retail", "SMB",
               "Specialty Retail", "Starbucks", "TJX", "TJX-FEDEX","Vouchering", "Walmart", "Weekend Coverage", "Whole Foods"]

@app.route("/it_form", methods =["GET", "POST"])
def it_form():
    if request.method == "POST":
        values = []
        errors = []
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        ticket_no = generate_ticket_no()
        values.extend([timestamp, ticket_no])

        previous_values = {}

        for field in IT_FIELDS:
            val = request.form.get(field, "").strip()
            previous_values[field] = val  # Store all user input
            if not val:
                errors.append(f"Please complete the '{field}' field.")
            values.append(val)

        # If there are errors, flash all of them and re-render the form
        if errors:
            for error in errors:
                flash(error, "danger")
            return render_template(
                "it_form.html",
                it_fields=IT_FIELDS,
                it_multiline_fields=IT_MULTILINE_FIELDS,
                today=datetime.now().strftime("%Y-%m-%d"),
                location_options=LOCATION_OPTIONS,
                request_options=REQUEST_OPTIONS,
                lob_options=LOB_OPTIONS,
                previous_values=previous_values  # So fields retain user input
            )

        # Email validation
        if not is_valid_email(values[2]):
            flash("Invalid email address.", "danger")
            today_str = datetime.now().strftime("%Y-%m-%d")
            return render_template(
                "it_form.html",
                it_fields=IT_FIELDS,
                it_multiline_fields=IT_MULTILINE_FIELDS,
                today=today_str,
                location_options=LOCATION_OPTIONS,
                request_options=REQUEST_OPTIONS,
                lob_options=LOB_OPTIONS,
                previous_values=request.form
            )

        # Save to Google Sheet
        try:
            it_sheet.append_row(values)
        except Exception as e:
            flash("An error occurred while submitting to the spreadsheet. Please try again later.", "danger")
            print(f"[GSheet Error] {e}")
            return redirect(request.url)

        # Email
        # Determine recipients
        recipients = [values[2], "manzonmarkjerwin@gmail.com"]

        #  Eliminate duplicates
        recipients = list(set(email.lower() for email in recipients))

        send_email(
            recipients,
            f"Automated Notification: IT Ticket Sent & Logged – {ticket_no}",
            f"""This is an automated notification to inform you that a IT Ticket has been sent and logged in to the system:
    
Ticket no. {ticket_no}
    
The matter is currently under review. Updates will be provided as they become available.
    
This is a system-generated message. No action is required unless otherwise indicated.""",
        )

        flash("IT Ticket submitted successfully.", "success")
        return redirect(f"/it_form?generated={ticket_no}")

    # GET method (default view)
    today_str = datetime.now().strftime("%Y-%m-%d")
    return render_template(
        "it_form.html",
        it_fields=IT_FIELDS,
        it_multiline_fields=IT_MULTILINE_FIELDS,
        today=today_str,
        location_options=LOCATION_OPTIONS,
        request_options=REQUEST_OPTIONS,
        lob_options=LOB_OPTIONS
    )
if __name__ == "__main__":
    app.run(debug=True)